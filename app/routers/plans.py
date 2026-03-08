import logging
import threading
import uuid

import httpx
from fastapi import APIRouter, Depends, HTTPException, Query, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.dependencies import get_current_user, get_db_session
from app.models.plan import DevelopmentPlan, SubmissionDocument
from app.schemas.common import JobAccepted
from app.schemas.plan import (
    ContractorRecommendationsResponse,
    ContractorResult,
    PlanClarifyResponse,
    PlanGenerateRequest,
    PlanListResponse,
    PlanResponse,
    PlanSubmissionReadinessResponse,
    ReviewActionRequest,
    SubmissionDocumentResponse,
)
from app.services.contractor_trades import derive_trade_categories
from app.services.submission.readiness import evaluate_submission_readiness
from app.services.submission.review import (
    approve_document,
    reject_document,
    submit_for_review,
)
from app.tasks.plan import run_plan_generation

router = APIRouter()


@router.post("/plans/generate", response_model=JobAccepted, status_code=status.HTTP_202_ACCEPTED)
async def generate_plan(
    body: PlanGenerateRequest,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    """Submit a natural language development query to generate a full plan.

    The system will:
    1. Parse your query into structured development parameters
    2. Look up the parcel and applicable zoning
    3. Generate building massing within policy constraints
    4. Optimize the unit mix
    5. Run financial pro forma
    6. Check entitlement compliance
    7. Search for precedent applications
    8. Generate government submission documents
    """
    plan = DevelopmentPlan(
        organization_id=user["organization_id"],
        created_by=user["id"],
        original_query=body.query,
        status="pending",
    )
    db.add(plan)
    await db.flush()
    await db.refresh(plan)
    await db.commit()

    threading.Thread(
        target=run_plan_generation,
        args=(str(plan.id), body.query, body.auto_run, body.generate_subset),
        daemon=True,
    ).start()

    return JobAccepted(
        job_id=plan.id,
        status="accepted",
        location=f"{settings.API_V1_PREFIX}/plans/{plan.id}",
    )


@router.get("/plans", response_model=list[PlanListResponse])
async def list_plans(
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    result = await db.execute(
        select(DevelopmentPlan)
        .where(DevelopmentPlan.organization_id == user["organization_id"])
        .order_by(DevelopmentPlan.created_at.desc())
    )
    return result.scalars().all()


@router.get("/plans/{plan_id}", response_model=PlanResponse)
async def get_plan(
    plan_id: uuid.UUID,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    result = await db.execute(
        select(DevelopmentPlan)
        .options(selectinload(DevelopmentPlan.documents))
        .where(
            DevelopmentPlan.id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(status_code=404, detail="Plan not found")
    return plan


@router.get("/plans/{plan_id}/readiness", response_model=PlanSubmissionReadinessResponse)
async def get_plan_readiness(
    plan_id: uuid.UUID,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    result = await db.execute(
        select(DevelopmentPlan)
        .options(selectinload(DevelopmentPlan.documents))
        .where(
            DevelopmentPlan.id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(status_code=404, detail="Plan not found")
    return evaluate_submission_readiness(plan, plan.documents)


@router.post("/plans/{plan_id}/clarify", response_model=PlanResponse)
async def clarify_plan(
    plan_id: uuid.UUID,
    body: PlanClarifyResponse,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    """Provide answers to clarification questions and resume pipeline."""
    result = await db.execute(
        select(DevelopmentPlan)
        .options(selectinload(DevelopmentPlan.documents))
        .where(
            DevelopmentPlan.id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(status_code=404, detail="Plan not found")
    if plan.status != "needs_clarification":
        raise HTTPException(status_code=400, detail=f"Plan is not awaiting clarification (status: {plan.status})")

    # Merge answers into parsed parameters
    params = plan.parsed_parameters or {}
    params["user_clarifications"] = body.answers
    plan.parsed_parameters = params
    plan.status = "parsed"
    await db.flush()
    await db.refresh(plan)
    await db.commit()

    # Resume pipeline
    threading.Thread(
        target=run_plan_generation,
        args=(str(plan.id), plan.original_query, True),
        daemon=True,
    ).start()

    return plan


@router.get("/plans/{plan_id}/documents", response_model=list[SubmissionDocumentResponse])
async def list_plan_documents(
    plan_id: uuid.UUID,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    result = await db.execute(
        select(SubmissionDocument)
        .join(DevelopmentPlan)
        .where(
            SubmissionDocument.plan_id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
        .order_by(SubmissionDocument.sort_order)
    )
    return result.scalars().all()


@router.get("/plans/{plan_id}/documents/{doc_id}", response_model=SubmissionDocumentResponse)
async def get_plan_document(
    plan_id: uuid.UUID,
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    result = await db.execute(
        select(SubmissionDocument)
        .join(DevelopmentPlan)
        .where(
            SubmissionDocument.id == doc_id,
            SubmissionDocument.plan_id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.post(
    "/plans/{plan_id}/documents/{doc_id}/submit-review",
    response_model=SubmissionDocumentResponse,
)
async def submit_document_for_review(
    plan_id: uuid.UUID,
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    """Submit a generated document for human review."""
    # Verify ownership
    plan_result = await db.execute(
        select(DevelopmentPlan).where(
            DevelopmentPlan.id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    if not plan_result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="Plan not found")

    try:
        doc = await submit_for_review(db, doc_id, user["id"])
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return doc


@router.post(
    "/plans/{plan_id}/documents/{doc_id}/approve",
    response_model=SubmissionDocumentResponse,
)
async def approve_plan_document(
    plan_id: uuid.UUID,
    doc_id: uuid.UUID,
    body: ReviewActionRequest,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    """Approve a document after review."""
    plan_result = await db.execute(
        select(DevelopmentPlan).where(
            DevelopmentPlan.id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    if not plan_result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="Plan not found")

    try:
        doc = await approve_document(db, doc_id, user["id"], body.notes)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return doc


@router.post(
    "/plans/{plan_id}/documents/{doc_id}/reject",
    response_model=SubmissionDocumentResponse,
)
async def reject_plan_document(
    plan_id: uuid.UUID,
    doc_id: uuid.UUID,
    body: ReviewActionRequest,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    """Reject a document — returns it to draft for revision."""
    plan_result = await db.execute(
        select(DevelopmentPlan).where(
            DevelopmentPlan.id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    if not plan_result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="Plan not found")

    try:
        doc = await reject_document(db, doc_id, user["id"], body.notes)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return doc


logger = logging.getLogger(__name__)

# ─── Valid doc types for regeneration ───
_VALID_DOC_TYPES = {
    "cover_letter", "planning_rationale", "compliance_matrix", "site_plan_data",
    "massing_summary", "unit_mix_summary", "financial_feasibility", "precedent_report",
    "public_benefit_statement", "shadow_study", "four_statutory_tests",
    "approval_pathway_document", "due_diligence_report", "olt_appeal_brief",
    "revised_rationale", "mediation_strategy", "neighbour_support_letter",
    "pac_prep_package", "submission_readiness_report", "correction_response",
    "compliance_review_report", "variance_justification", "as_of_right_check",
    "required_studies_checklist", "timeline_cost_estimate",
    "building_permit_readiness_checklist", "professional_referral_checklist",
}

_RULE_BASED_DOC_TYPES = {
    "as_of_right_check", "required_studies_checklist", "timeline_cost_estimate",
    "building_permit_readiness_checklist", "professional_referral_checklist",
}


@router.post(
    "/plans/{plan_id}/generate-document/{doc_type}",
    response_model=SubmissionDocumentResponse,
)
async def regenerate_document(
    plan_id: uuid.UUID,
    doc_type: str,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    """Generate or regenerate a single document from an existing completed plan."""
    from app.ai.factory import get_ai_provider
    from app.schemas.plan import PlanGenerateDocumentRequest
    from app.services.submission.context_builder import build_document_context
    from app.services.submission.generator import SubmissionPackageGenerator
    from app.tasks.plan import SUBMISSION_DOCUMENTS

    if doc_type not in _VALID_DOC_TYPES:
        raise HTTPException(
            status_code=422,
            detail=f"Invalid doc_type: {doc_type}. Must be one of: {sorted(_VALID_DOC_TYPES)}",
        )

    # Load plan
    result = await db.execute(
        select(DevelopmentPlan)
        .options(selectinload(DevelopmentPlan.documents))
        .where(
            DevelopmentPlan.id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(status_code=404, detail="Plan not found")
    if plan.status != "completed":
        raise HTTPException(status_code=400, detail=f"Plan must be completed (current: {plan.status})")

    # Rebuild context from stored plan data
    summary = plan.summary or {}
    parsed = plan.parsed_parameters or {}
    context = build_document_context(
        parcel_data={
            "address": parsed.get("resolved_address") or parsed.get("address"),
            "zone_code": parsed.get("zone_code"),
            "lot_area_m2": None,
            "lot_frontage_m": None,
            "lot_depth_m": None,
            "current_use": None,
        },
        zoning=None,
        massing=summary.get("massing"),
        layout=summary.get("layout"),
        finance=summary.get("finance"),
        compliance=None,
        precedents=None,
        policy_stack=None,
        overlays=None,
        project_name=parsed.get("project_name", ""),
        organization_name="",
        parsed_parameters=parsed,
    )

    # Generate content
    provider = get_ai_provider()
    generator = SubmissionPackageGenerator(provider)

    if doc_type in _RULE_BASED_DOC_TYPES:
        gen_result = generator.generate_rule_based_document(doc_type, context)
    else:
        gen_result = await generator.generate_document(doc_type, context)

    content = gen_result["content_text"]
    content_json = gen_result.get("content_json")
    metadata = gen_result.get("metadata", {})

    # Find doc spec for title/description
    doc_spec = next((d for d in SUBMISSION_DOCUMENTS if d["doc_type"] == doc_type), None)
    title = doc_spec["title"] if doc_spec else doc_type.replace("_", " ").title()
    description = doc_spec.get("description", "") if doc_spec else ""
    sort_order = doc_spec.get("sort_order", 0) if doc_spec else 0

    # Upsert — update existing or create new
    existing = next((d for d in plan.documents if d.doc_type == doc_type), None)
    if existing:
        existing.content_text = content
        existing.content_json = content_json
        existing.status = "completed"
        existing.review_status = "draft"
        existing.ai_provider = metadata.get("ai_provider")
        existing.ai_model = metadata.get("ai_model")
        doc = existing
    else:
        doc = SubmissionDocument(
            plan_id=plan.id,
            doc_type=doc_type,
            title=title,
            description=description,
            sort_order=sort_order,
            format="markdown",
            status="completed",
            review_status="draft",
            content_text=content,
            content_json=content_json,
            ai_provider=metadata.get("ai_provider"),
            ai_model=metadata.get("ai_model"),
        )
        db.add(doc)

    await db.flush()
    await db.refresh(doc)
    await db.commit()
    return doc


@router.get("/plans/{plan_id}/documents/{doc_id}/download")
async def download_document(
    plan_id: uuid.UUID,
    doc_id: uuid.UUID,
    format: str = Query(default="markdown", regex="^(markdown|html|docx)$"),
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    """Download a document in the specified format."""
    from fastapi.responses import StreamingResponse
    import io

    result = await db.execute(
        select(SubmissionDocument)
        .join(DevelopmentPlan)
        .where(
            SubmissionDocument.id == doc_id,
            SubmissionDocument.plan_id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    content = doc.content_text or ""
    filename_base = f"{doc.doc_type}_{doc.title.replace(' ', '_')}"

    if format == "markdown":
        return StreamingResponse(
            io.BytesIO(content.encode("utf-8")),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.md"'},
        )
    elif format == "html":
        try:
            import markdown as md_lib
            html_content = md_lib.markdown(content, extensions=["tables"])
        except ImportError:
            html_content = f"<pre>{content}</pre>"
        styled_html = (
            "<!DOCTYPE html><html><head><meta charset='utf-8'>"
            "<style>body{font-family:sans-serif;max-width:800px;margin:40px auto;padding:0 20px;}"
            "table{border-collapse:collapse;width:100%;}th,td{border:1px solid #ddd;padding:8px;text-align:left;}"
            "th{background:#f5f5f5;}blockquote{border-left:4px solid #ffc107;padding:10px;background:#fffde7;}"
            "</style></head><body>" + html_content + "</body></html>"
        )
        return StreamingResponse(
            io.BytesIO(styled_html.encode("utf-8")),
            media_type="text/html",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.html"'},
        )
    else:  # docx
        try:
            from docx import Document as DocxDocument
            docx_doc = DocxDocument()
            for line in content.split("\n"):
                stripped = line.strip()
                if stripped.startswith("# "):
                    docx_doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    docx_doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    docx_doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("> "):
                    p = docx_doc.add_paragraph(stripped[2:])
                    p.style = "Intense Quote"
                elif stripped.startswith("- "):
                    docx_doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped:
                    docx_doc.add_paragraph(stripped)
            buf = io.BytesIO()
            docx_doc.save(buf)
            buf.seek(0)
            return StreamingResponse(
                buf,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename_base}.docx"'},
            )
        except ImportError:
            raise HTTPException(status_code=501, detail="python-docx not installed")


@router.post("/plans/{plan_id}/export")
async def export_plan(
    plan_id: uuid.UUID,
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    """Export all plan documents as a single DOCX file."""
    from fastapi.responses import StreamingResponse
    import io

    result = await db.execute(
        select(DevelopmentPlan)
        .options(selectinload(DevelopmentPlan.documents))
        .where(
            DevelopmentPlan.id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(status_code=404, detail="Plan not found")

    docs = sorted(plan.documents, key=lambda d: d.sort_order or 0)
    if not docs:
        raise HTTPException(status_code=400, detail="No documents to export")

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
        from docx.enum.section import WD_ORIENT

        docx_doc = DocxDocument()
        style = docx_doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Calibri"

        for idx, doc in enumerate(docs):
            if idx > 0:
                docx_doc.add_page_break()

            docx_doc.add_heading(doc.title or doc.doc_type.replace("_", " ").title(), level=0)

            content = doc.content_text or ""
            for line in content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("# "):
                    docx_doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    docx_doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    docx_doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("#### "):
                    docx_doc.add_heading(stripped[5:], level=4)
                elif stripped.startswith("> "):
                    p = docx_doc.add_paragraph(stripped[2:])
                    p.style = "Intense Quote"
                elif stripped.startswith("- "):
                    docx_doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. "):
                    docx_doc.add_paragraph(stripped[3:], style="List Number")
                elif stripped:
                    docx_doc.add_paragraph(stripped)

        buf = io.BytesIO()
        docx_doc.save(buf)
        buf.seek(0)

        plan_name = (plan.parsed_parameters or {}).get("project_name", "submission_package")
        filename = f"{plan_name.replace(' ', '_')}_export.docx"

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx not installed")


@router.get("/plans/{plan_id}/contractors", response_model=ContractorRecommendationsResponse)
async def get_contractor_recommendations(
    plan_id: uuid.UUID,
    lat: float = Query(..., description="Latitude for nearby search"),
    lng: float = Query(..., description="Longitude for nearby search"),
    db: AsyncSession = Depends(get_db_session),
    user: dict = Depends(get_current_user),
):
    """Recommend local contractors based on plan outputs."""
    if not settings.GOOGLE_PLACES_API_KEY:
        return ContractorRecommendationsResponse(contractors=[])

    result = await db.execute(
        select(DevelopmentPlan)
        .options(selectinload(DevelopmentPlan.documents))
        .where(
            DevelopmentPlan.id == plan_id,
            DevelopmentPlan.organization_id == user["organization_id"],
        )
    )
    plan = result.scalar_one_or_none()
    if not plan:
        raise HTTPException(status_code=404, detail="Plan not found")

    doc_types = [d.doc_type for d in plan.documents]
    massing = (plan.summary or {}).get("massing")
    trades = derive_trade_categories(doc_types, massing)

    contractors: list[ContractorResult] = []
    async with httpx.AsyncClient(timeout=10) as client:
        for trade in trades:
            try:
                resp = await client.get(
                    "https://maps.googleapis.com/maps/api/place/textsearch/json",
                    params={
                        "query": f"{trade} near Toronto",
                        "location": f"{lat},{lng}",
                        "radius": "10000",
                        "key": settings.GOOGLE_PLACES_API_KEY,
                    },
                )
                data = resp.json()
                for place in (data.get("results") or [])[:3]:
                    contractors.append(ContractorResult(
                        name=place.get("name", ""),
                        rating=place.get("rating"),
                        review_count=place.get("user_ratings_total"),
                        phone=place.get("formatted_phone_number"),
                        website=None,
                        address=place.get("formatted_address"),
                        trade=trade,
                    ))
            except Exception:
                logger.warning("Google Places API error for trade=%s", trade, exc_info=True)
                continue

    return ContractorRecommendationsResponse(contractors=contractors)
